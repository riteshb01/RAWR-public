"""
Bengal RAWR — Document Parser
Extracts clean text from PDF, TXT, and DOCX files.
Supports OCR fallback via pytesseract for scanned PDFs.
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def clean_extracted_text(text: str) -> str:
    """
    Normalize extracted text for NLP processing.
    - Collapse whitespace
    - Remove page numbers
    - Normalize line endings
    - Remove headers/footers heuristically
    """
    if not text:
        return ''

    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Remove page number patterns (standalone numbers or "Page X of Y")
    text = re.sub(r'^\s*Page\s+\d+\s+of\s+\d+\s*$', '', text, flags=re.MULTILINE | re.IGNORECASE)
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)

    # Remove lines that are only dashes, underscores or asterisks (decorative)
    text = re.sub(r'^\s*[-_*=]{3,}\s*$', '', text, flags=re.MULTILINE)

    # Collapse multiple blank lines into a single blank line
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove excessive leading/trailing whitespace per line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)

    return text.strip()


def extract_pdf_text(file_path: str) -> str:
    """
    Extract text from a PDF using pdfplumber.
    Falls back to pytesseract OCR if a page has no extractable text
    (handles scanned/image-based PDFs).
    """
    text_parts = []

    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required: pip install pdfplumber")

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()

                if page_text and page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"PDF page {page_num}: extracted {len(page_text)} chars")
                else:
                    # Fallback to OCR for image-based pages
                    ocr_text = _ocr_page(page, page_num)
                    if ocr_text:
                        text_parts.append(ocr_text)
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}")
        raise

    full_text = '\n\n'.join(text_parts)
    return clean_extracted_text(full_text)


def _ocr_page(page, page_num: int) -> Optional[str]:
    """OCR a single PDF page using pytesseract."""
    try:
        import pytesseract
        from PIL import Image
        import io

        # Render page to image at 200 DPI
        img = page.to_image(resolution=200).original
        ocr_text = pytesseract.image_to_string(img)
        logger.info(f"OCR fallback used on page {page_num}")
        return ocr_text
    except ImportError:
        logger.warning("pytesseract/PIL not installed — skipping OCR fallback")
        return None
    except Exception as e:
        logger.warning(f"OCR failed on page {page_num}: {e}")
        return None


def extract_txt_text(file_path: str) -> str:
    """
    Extract text from a plain text file.
    Tries common encodings.
    """
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
            logger.debug(f"TXT file read with encoding {encoding}")
            return clean_extracted_text(text)
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise

    raise ValueError(f"Could not decode text file {file_path} with any supported encoding")


def extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    try:
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = '\n'.join(paragraphs)
        return clean_extracted_text(full_text)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {e}")
        raise


class DocumentParser:
    """
    Unified document parser that routes to the correct extractor
    based on file extension.
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx', '.doc'}

    def __init__(self):
        self.supported_extensions = self.SUPPORTED_EXTENSIONS

    def parse(self, file_path: str) -> str:
        """
        Parse a document and return clean text.

        Args:
            file_path: Absolute path to the uploaded file

        Returns:
            Cleaned text string

        Raises:
            ValueError: If file type is unsupported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"Parsing {ext.upper()} document: {path.name}")

        if ext == '.pdf':
            return extract_pdf_text(file_path)
        elif ext == '.txt':
            return extract_txt_text(file_path)
        elif ext in ('.docx', '.doc'):
            return extract_docx_text(file_path)

    def get_file_metadata(self, file_path: str) -> dict:
        """Return file metadata: size, extension, name."""
        path = Path(file_path)
        return {
            'filename': path.name,
            'extension': path.suffix.lower(),
            'size_bytes': path.stat().st_size if path.exists() else 0,
            'size_kb': round(path.stat().st_size / 1024, 2) if path.exists() else 0,
        }
